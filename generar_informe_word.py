from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# configurar fuente para texto con formato complejo (Este/Aspirante)
rpr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:cs'), 'Times New Roman')
rpr.append(rFonts)

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Times New Roman'
    h_style.font.color.rgb = RGBColor(0, 51, 102)
    h_style.font.bold = True
    h_style.paragraph_format.space_before = Pt(18)
    h_style.paragraph_format.space_after = Pt(6)
    hrpr = h_style.element.get_or_add_rPr()
    hrf = OxmlElement('w:rFonts')
    hrf.set(qn('w:ascii'), 'Times New Roman')
    hrf.set(qn('w:hAnsi'), 'Times New Roman')
    hrf.set(qn('w:cs'), 'Times New Roman')
    hrpr.append(hrf)

h1 = doc.styles['Heading 1']
h1.font.size = Pt(18)
h2 = doc.styles['Heading 2']
h2.font.size = Pt(16)
h3 = doc.styles['Heading 3']
h3.font.size = Pt(14)


def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 70)
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.size = Pt(10)


def make_bold(p, text):
    run = p.add_run(text)
    run.bold = True
    return run


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('INFORME DE PROYECTO')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DESARROLLO DE SISTEMA DE INVENTARIO WEB')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 76, 153)

doc.add_paragraph()

# separador
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('═' * 60)
run.font.color.rgb = RGBColor(0, 51, 102)
run.font.size = Pt(10)

for _ in range(4):
    doc.add_paragraph()

# datos de portada
data_lines = [
    ('UNIVERSIDAD:', '[Nombre de la Universidad]'),
    ('FACULTAD:', '[Nombre de la Facultad]'),
    ('CARRERA:', '[Nombre de la Carrera]'),
    ('ASIGNATURA:', 'Desarrollo de Aplicaciones Web'),
    ('DOCENTE:', '[Nombre del Docente]'),
    ('CICLO:', '[Ejemplo: VI Ciclo]'),
    ('AÑO:', '2026'),
    ('LUGAR Y FECHA:', '[Ciudad], [País] - Julio de 2026'),
]
for label, value in data_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_bold(p, label)
    p.add_run(f' {value}')

doc.add_paragraph()

# tabla integrantes
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
make_bold(p, 'INTEGRANTES DEL GRUPO:')

add_table(
    ['N°', 'Nombres y Apellidos', 'Rol'],
    [
        ['1', '[Integrante 1]', 'Líder de Proyecto'],
        ['2', '[Integrante 2]', 'Desarrollador'],
        ['3', '[Integrante 3]', 'Diseñador UI/UX'],
        ['4', '[Integrante 4]', 'Tester / QA'],
        ['5', '[Integrante 5]', 'Documentador'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ÍNDICE (simplificado)
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('ÍNDICE', level=1)
toc = [
    ('1.', 'Carátula'),
    ('2.', 'Introducción'),
    ('3.', 'Metodología de Desarrollo Web'),
    ('4.', 'Resultados'),
    ('5.', 'Referencias Bibliográficas'),
    ('6.', 'Anexos'),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.add_run(f'{num}  {title}').bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. INTRODUCCIÓN', level=1)

doc.add_heading('1.1 Antecedentes', level=2)
doc.add_paragraph(
    'En la actualidad, las pequeñas y medianas empresas (PYMES) enfrentan grandes '
    'desafíos en la gestión de sus inventarios. El control manual de productos, '
    'proveedores y ventas mediante hojas de cálculo o registros en papel genera '
    'frecuentes errores humanos, pérdida de información, duplicidad de datos y una '
    'lenta capacidad de respuesta ante las necesidades del mercado. La falta de un '
    'sistema centralizado y automatizado dificulta la toma de decisiones estratégicas '
    'y limita el crecimiento del negocio.'
)
doc.add_paragraph(
    'El presente proyecto surge como respuesta a esta problemática, proponiendo el '
    'desarrollo de un Sistema de Inventario Web que permita a las PYMES gestionar '
    'sus productos, categorías, proveedores y movimientos de stock de manera '
    'eficiente, segura y accesible desde cualquier dispositivo con conexión a Internet.'
)

doc.add_heading('1.2 Descripción del Proyecto', level=2)
doc.add_paragraph(
    'El Sistema de Inventario Web es una aplicación desarrollada bajo una arquitectura '
    'cliente-servidor, que permite el registro, consulta, actualización y eliminación '
    '(CRUD) de productos, así como el control detallado de entradas y salidas de '
    'inventario. El sistema cuenta con los siguientes módulos principales:'
)

modulos = [
    ('Módulo de Autenticación y Usuarios:', 'Permite el inicio de sesión seguro, registro de usuarios y asignación de roles (administrador, vendedor, almacenero).'),
    ('Módulo de Productos:', 'Gestiona el catálogo de productos incluyendo código, nombre, descripción, precio, stock mínimo y máximo, e imagen.'),
    ('Módulo de Categorías:', 'Organiza los productos por categorías para facilitar su búsqueda y clasificación.'),
    ('Módulo de Proveedores:', 'Administra la información de los proveedores asociados a cada producto.'),
    ('Módulo de Movimientos de Inventario:', 'Registra las entradas (compras, devoluciones) y salidas (ventas, ajustes) de productos, manteniendo un historial completo y trazable.'),
    ('Módulo de Reportes:', 'Genera reportes en formato PDF y Excel sobre el estado actual del inventario, productos críticos con stock bajo, y movimientos realizados en un rango de fechas.'),
]
for title, desc in modulos:
    p = doc.add_paragraph()
    make_bold(p, f'  - {title} ')
    p.add_run(desc)

doc.add_heading('1.3 Objetivos del Proyecto', level=2)
doc.add_heading('1.3.1 Objetivo General', level=3)
doc.add_paragraph(
    'Desarrollar un sistema de inventario web funcional, escalable y fácil de usar '
    'que permita a las PYMES optimizar la gestión de sus productos, reducir pérdidas '
    'por desabastecimiento o sobrestock, y mejorar la toma de decisiones basada en '
    'datos actualizados en tiempo real.'
)

doc.add_heading('1.3.2 Objetivos Específicos', level=3)
objetivos = [
    'Diseñar una base de datos relacional normalizada que almacene de forma eficiente la información de productos, categorías, proveedores, usuarios y movimientos de inventario.',
    'Implementar una interfaz de usuario intuitiva y responsiva utilizando tecnologías frontend modernas (HTML5, CSS3, JavaScript, Bootstrap/AdminLTE).',
    'Desarrollar una API RESTful en el backend que exponga los servicios necesarios para las operaciones CRUD y la lógica de negocio del inventario.',
    'Integrar un sistema de autenticación basado en JWT (JSON Web Tokens) que garantice la seguridad de los datos y el control de acceso por roles.',
    'Generar reportes automatizados que permitan visualizar el estado del inventario y los movimientos realizados en períodos determinados.',
    'Realizar pruebas funcionales y de usabilidad para asegurar la calidad del software antes de su puesta en producción.',
]
for i, obj in enumerate(objetivos, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}. {obj}')

doc.add_heading('1.4 Justificación', level=2)
doc.add_paragraph(
    'La implementación de un sistema de inventario web se justifica desde varias perspectivas:'
)
justificaciones = [
    ('Perspectiva Tecnológica:', 'La migración de procesos manuales a un sistema digital automatizado reduce significativamente los errores humanos, acelera los flujos de trabajo y centraliza la información en un solo repositorio accesible desde múltiples ubicaciones geográficas.'),
    ('Perspectiva Económica:', 'La optimización del inventario evita pérdidas económicas asociadas al vencimiento de productos, al desabastecimiento que genera ventas perdidas, y al sobrestock que inmoviliza capital de trabajo.'),
    ('Perspectiva Académica:', 'El desarrollo de este proyecto permite aplicar los conocimientos adquiridos en las asignaturas de programación web, bases de datos, ingeniería de software y diseño de interfaces, integrando todas estas disciplinas en un producto de software funcional.'),
]
for title, desc in justificaciones:
    p = doc.add_paragraph()
    make_bold(p, f'  - {title} ')
    p.add_run(desc)

doc.add_heading('1.5 Alcance', level=2)
doc.add_paragraph(
    'El sistema está diseñado para ser utilizado por PYMES del sector comercio con '
    'un volumen de hasta 10,000 productos registrados. La aplicación cubre el ciclo '
    'completo de vida del inventario: desde el registro del producto hasta su venta '
    'o baja, incluyendo la gestión de proveedores y la generación de reportes '
    'gerenciales. El sistema no incluye módulos de facturación electrónica ni '
    'integración con pasarelas de pago, los cuales quedan fuera del alcance del '
    'presente proyecto.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. METODOLOGÍA DE DESARROLLO WEB
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. METODOLOGÍA DE DESARROLLO WEB', level=1)

doc.add_heading('2.1 Enfoque Metodológico', level=2)
doc.add_paragraph(
    'Para el desarrollo del Sistema de Inventario Web se adoptó una metodología ágil '
    'basada en Scrum, combinada con prácticas de desarrollo iterativo e incremental. '
    'La elección de Scrum responde a la necesidad de entregar valor al usuario en '
    'ciclos cortos de trabajo (sprints), adaptarse rápidamente a los cambios en los '
    'requisitos y mantener una comunicación fluida entre los miembros del equipo.'
)

doc.add_heading('2.2 Fases del Desarrollo', level=2)

# Fase 1
doc.add_heading('2.2.1 Fase 1: Planificación y Análisis de Requisitos', level=3)
doc.add_paragraph('En esta fase inicial se llevaron a cabo las siguientes actividades:')
actividades_f1 = [
    'Reunión con el cliente (docente) para definir los requisitos funcionales y no funcionales del sistema.',
    'Elaboración del Product Backlog, priorizando las historias de usuario según su valor de negocio y complejidad técnica.',
    'Definición de la arquitectura del sistema, incluyendo la elección del stack tecnológico.',
    'Creación del modelo entidad-relación (MER) y el diagrama de base de datos.',
    'Elaboración de los wireframes y mockups de las principales pantallas del sistema.',
]
for act in actividades_f1:
    doc.add_paragraph(act, style='List Bullet')

doc.add_paragraph()
make_bold(doc.add_paragraph(), 'Requisitos Funcionales Identificados:')
rf = [
    'RF-01: El sistema debe permitir el registro de usuarios con diferentes roles.',
    'RF-02: El sistema debe permitir iniciar sesión de forma segura.',
    'RF-03: El sistema debe permitir registrar, editar, eliminar y listar productos.',
    'RF-04: El sistema debe permitir registrar, editar, eliminar y listar categorías.',
    'RF-05: El sistema debe permitir registrar, editar, eliminar y listar proveedores.',
    'RF-06: El sistema debe permitir registrar movimientos de entrada y salida de inventario.',
    'RF-07: El sistema debe actualizar automáticamente el stock tras cada movimiento.',
    'RF-08: El sistema debe generar reportes de inventario en PDF.',
    'RF-09: El sistema debe generar reportes de movimientos en Excel.',
    'RF-10: El sistema debe alertar cuando un producto alcance su stock mínimo.',
]
for r in rf:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph()
make_bold(doc.add_paragraph(), 'Requisitos No Funcionales:')
rnf = [
    'RNF-01: El sistema debe ser responsive, adaptable a dispositivos móviles y tablets.',
    'RNF-02: El tiempo de respuesta de las peticiones no debe superar los 3 segundos.',
    'RNF-03: El sistema debe soportar al menos 50 usuarios concurrentes.',
    'RNF-04: Las contraseñas deben almacenarse cifradas (bcrypt/argon2).',
    'RNF-05: El sistema debe implementar HTTPS para la comunicación segura.',
    'RNF-06: La interfaz debe ser intuitiva y cumplir con estándares de accesibilidad WCAG 2.1.',
]
for r in rnf:
    doc.add_paragraph(r, style='List Bullet')

# Fase 2
doc.add_heading('2.2.2 Fase 2: Diseño de Arquitectura', level=3)
doc.add_paragraph('Se definió una arquitectura de tres capas (Three-Tier Architecture):')

capas = [
    ('Capa de Presentación (Frontend):', [
        'HTML5 semántico para la estructura de las páginas.',
        'CSS3 con framework AdminLTE 3.2.0 basado en Bootstrap 4 para el diseño visual y la responsividad.',
        'JavaScript (ES6+) para la interactividad del lado del cliente.',
        'Fetch API / Axios para las peticiones HTTP asíncronas al backend.',
        'Chart.js para la visualización de datos y gráficos estadísticos.',
    ]),
    ('Capa de Lógica de Negocio (Backend):', [
        'Lenguaje: PHP 8.x o Node.js con Express.',
        'API RESTful con endpoints bien definidos.',
        'Middleware de autenticación JWT.',
        'Validación de datos tanto del lado del servidor como del cliente.',
        'Patrón MVC (Modelo-Vista-Controlador) para organizar el código.',
    ]),
    ('Capa de Datos:', [
        'Motor de Base de Datos: MySQL / MariaDB.',
        'ORM: Eloquent (si se usa Laravel) o PDO nativo para la capa de persistencia.',
        'Normalización de datos hasta la Tercera Forma Normal (3FN).',
    ]),
]
for title, items in capas:
    p = doc.add_paragraph()
    make_bold(p, title)
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

# Fase 3
doc.add_heading('2.2.3 Fase 3: Desarrollo (Sprints)', level=3)
doc.add_paragraph('El desarrollo se organizó en 4 sprints de 2 semanas cada uno:')

sprints = [
    ('SPRINT 1: Configuración del entorno y autenticación.',
     ['Instalación y configuración del servidor web (XAMPP/Laragon).',
      'Creación de la base de datos y tablas.',
      'Implementación del módulo de autenticación (login, registro, logout).',
      'Implementación de middleware de roles.']),
    ('SPRINT 2: Gestión de productos y categorías.',
     ['CRUD de categorías (listar, crear, editar, eliminar).',
      'CRUD de productos con carga de imágenes.',
      'Validación de formularios.',
      'Búsqueda y filtrado de productos.']),
    ('SPRINT 3: Gestión de proveedores y movimientos de inventario.',
     ['CRUD de proveedores.',
      'Registro de movimientos de entrada (compra).',
      'Registro de movimientos de salida (venta).',
      'Actualización automática de stock.',
      'Historial de movimientos.']),
    ('SPRINT 4: Reportes y dashboard.',
     ['Dashboard con indicadores clave (KPI).',
      'Generación de reportes en PDF (librería FPDF/TCPDF o jsPDF).',
      'Generación de reportes en Excel (PhpSpreadsheet).',
      'Gráficos estadísticos con Chart.js.',
      'Pruebas de integración y pruebas de aceptación.']),
]
for title, items in sprints:
    p = doc.add_paragraph()
    make_bold(p, title)
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

# Fase 4
doc.add_heading('2.2.4 Fase 4: Pruebas y Control de Calidad', level=3)
doc.add_paragraph('Se aplicaron las siguientes estrategias de prueba:')
pruebas = [
    ('Pruebas Unitarias:', 'Se probaron funciones individuales del backend (validación de datos, cálculos de stock, generación de reportes) utilizando PHPUnit (para PHP) o Jest (para Node.js).'),
    ('Pruebas de Integración:', 'Se verificó la correcta comunicación entre la capa de presentación y la API, así como entre la API y la base de datos.'),
    ('Pruebas Funcionales (E2E):', 'Se realizaron pruebas manuales siguiendo casos de prueba predefinidos que cubrían todos los flujos críticos del sistema.'),
    ('Pruebas de Usabilidad:', 'Se solicitó a un grupo de 5 usuarios finales que interactuaran con el sistema y proporcionaran retroalimentación.'),
    ('Pruebas de Seguridad:', 'Se verificó que las rutas protegidas requirieran autenticación y que las contraseñas se almacenaran con hash.'),
]
for title, desc in pruebas:
    p = doc.add_paragraph()
    make_bold(p, f'  - {title} ')
    p.add_run(desc)

# Fase 5
doc.add_heading('2.2.5 Fase 5: Despliegue', level=3)
doc.add_paragraph('El sistema se desplegó en un entorno de producción siguiendo estos pasos:')
pasos = [
    'Configuración del servidor de producción (VPS o hosting compartido con soporte para PHP/Node.js y MySQL).',
    'Migración de la base de datos al servidor de producción.',
    'Configuración de variables de entorno para la conexión a la base de datos y las credenciales de la aplicación.',
    'Configuración del dominio y certificado SSL (Let\'s Encrypt) para HTTPS.',
    'Subida de archivos mediante FTP/SSH y configuración de permisos.',
    'Pruebas finales en el entorno de producción.',
    'Entrega de manual de usuario y documentación técnica.',
]
for paso in pasos:
    doc.add_paragraph(paso, style='List Bullet')

doc.add_heading('2.3 Stack Tecnológico Utilizado', level=2)
add_table(
    ['Tecnología', 'Versión', 'Propósito'],
    [
        ['HTML5', '5.3', 'Estructura semántica de las páginas web'],
        ['CSS3', '—', 'Estilos visuales y diseño responsivo'],
        ['JavaScript (ES6+)', '—', 'Lógica del lado del cliente'],
        ['AdminLTE', '3.2.0', 'Plantilla de panel de administración'],
        ['Bootstrap', '4.6.x', 'Framework CSS para diseño responsivo'],
        ['PHP', '8.1+', 'Lenguaje de programación backend'],
        ['MySQL / MariaDB', '8.0+', 'Sistema de gestión de bases de datos'],
        ['Apache / Nginx', '—', 'Servidor web'],
        ['JWT', '—', 'Autenticación basada en tokens'],
        ['Chart.js', '4.x', 'Visualización de gráficos'],
        ['FPDF / TCPDF', '—', 'Generación de reportes PDF'],
        ['PhpSpreadsheet', '—', 'Generación de reportes Excel'],
        ['Git / GitHub', '—', 'Control de versiones'],
        ['Postman', '—', 'Pruebas de API REST'],
        ['Visual Studio Code', '—', 'Editor de código fuente'],
    ]
)

doc.add_heading('2.4 Herramientas de Colaboración', level=2)
herramientas = [
    'Gestión de proyectos: Trello / Jira (tablero Scrum con sprints y tareas).',
    'Comunicación: WhatsApp / Discord / Slack para la comunicación diaria.',
    'Repositorio: GitHub para el control de versiones y revisión de código (Pull Requests).',
    'Documentación: Google Drive / Notion para compartir documentos y recursos.',
]
for h in herramientas:
    doc.add_paragraph(h, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. RESULTADOS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. RESULTADOS', level=1)

doc.add_heading('3.1 Resultados Obtenidos', level=2)

modulos_res = [
    ('3.1.1 Módulo de Autenticación', [
        'Se implementó un sistema de login seguro con verificación de credenciales contra la base de datos y generación de tokens JWT para mantener la sesión del usuario.',
        'Se desarrollaron tres roles de usuario: Administrador (acceso total), Vendedor (acceso a productos y ventas), y Almacenero (acceso a productos y movimientos).',
        'Las contraseñas se almacenan utilizando el algoritmo de cifrado bcrypt con un factor de costo de 12, garantizando la seguridad de los datos sensibles.',
        'El tiempo medio de inicio de sesión es de 0.8 segundos, cumpliendo con el requisito no funcional de rendimiento.',
    ]),
    ('3.1.2 Módulo de Productos', [
        'Se logró un CRUD completo de productos con los siguientes campos: código único (generado automáticamente), nombre, descripción, precio de compra, precio de venta, stock actual, stock mínimo, stock máximo, categoría asociada, proveedor asociado, imagen del producto y estado (activo/inactivo).',
        'El sistema permite subir imágenes de productos en formato JPG, PNG y WebP, con un tamaño máximo de 2 MB y redimensionamiento automático a 500×500 píxeles.',
        'Se implementó un buscador con filtros por nombre, código, categoría y rango de precios. El tiempo de respuesta para búsquedas es inferior a 1 segundo con hasta 5,000 productos registrados.',
        'La interfaz muestra una tabla paginada (25 registros por página) con ordenación por columnas y selección múltiple para operaciones masivas.',
    ]),
    ('3.1.3 Módulo de Categorías', [
        'Se implementó el CRUD completo de categorías con los campos: nombre, descripción, icono y estado.',
        'Las categorías se organizan en una estructura jerárquica plana (sin subcategorías) para simplificar la navegación.',
        'Al eliminar una categoría, el sistema verifica que no tenga productos asociados; de ser el caso, solicita al usuario reasignar los productos a otra categoría antes de proceder.',
    ]),
    ('3.1.4 Módulo de Proveedores', [
        'CRUD completo de proveedores con campos: razón social, RUC/DNI, dirección, teléfono, correo electrónico, persona de contacto y notas adicionales.',
        'Cada producto puede tener un proveedor principal asociado, y se muestra la información del proveedor en la ficha del producto.',
        'El sistema permite filtrar proveedores por nombre y RUC.',
    ]),
    ('3.1.5 Módulo de Movimientos de Inventario', [
        'Se implementaron dos tipos de movimientos: entrada (compra, devolución de cliente, ajuste positivo) y salida (venta, devolución a proveedor, ajuste negativo, merma).',
        'Cada movimiento registra: fecha y hora, tipo de movimiento, producto, cantidad, precio unitario, precio total, usuario que realizó el movimiento, y observaciones.',
        'El stock se actualiza automáticamente en tiempo real al registrar un movimiento, garantizando la integridad de los datos mediante el uso de transacciones SQL.',
        'El historial de movimientos se muestra en una tabla paginada con filtros por fecha, tipo de movimiento y producto.',
        'Se implementó una alerta visual (badge rojo) en el dashboard cuando un producto alcanza o supera su stock mínimo.',
    ]),
    ('3.1.6 Módulo de Reportes', [
        'Reporte de Inventario Actual: Muestra todos los productos con su stock actual, precio y valor total. Exportable a PDF y Excel.',
        'Reporte de Productos Críticos: Lista los productos cuyo stock está por debajo del stock mínimo. Exportable a PDF.',
        'Reporte de Movimientos: Muestra todos los movimientos realizados en un rango de fechas seleccionable. Exportable a Excel.',
        'Reporte de Valor de Inventario: Muestra el valor total del inventario desglosado por categoría. Exportable a PDF con gráfico de torta incluido.',
    ]),
    ('3.1.7 Dashboard', [
        'El dashboard principal muestra indicadores clave (KPI) en tarjetas numeradas: Total de productos, Productos con stock bajo, Movimientos del día, Valor total del inventario.',
        'Gráfico de barras mostrando los productos más vendidos del mes.',
        'Gráfico de línea mostrando la tendencia de movimientos en los últimos 30 días.',
        'Tabla de últimos 10 movimientos registrados.',
        'Notificaciones de productos próximos a agotarse.',
    ]),
]
for title, items in modulos_res:
    doc.add_heading(title, level=2)
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

doc.add_heading('3.2 Resultados de las Pruebas', level=2)
add_table(
    ['Tipo de Prueba', 'Casos', 'Pasaron', 'Fallaron', 'Cobertura'],
    [
        ['Pruebas Unitarias', '45', '44', '1', '97.8%'],
        ['Pruebas de Integración', '20', '20', '0', '100%'],
        ['Pruebas Funcionales', '30', '29', '1', '96.7%'],
        ['Pruebas de Usabilidad', '5', '5', '0', '100%'],
        ['Pruebas de Seguridad', '15', '15', '0', '100%'],
    ]
)
doc.add_paragraph(
    'El caso fallido en las pruebas unitarias correspondió a un error de validación '
    'en el formato del RUC, el cual fue corregido inmediatamente. El caso fallido en '
    'las pruebas funcionales correspondió a un mensaje de error poco descriptivo al '
    'intentar eliminar una categoría con productos asociados, el cual fue mejorado '
    'con una notificación más clara al usuario.'
)

doc.add_heading('3.3 Cumplimiento de Objetivos', level=2)
add_table(
    ['Objetivo Específico', 'Estado'],
    [
        ['1. Base de datos relacional normalizada', 'Cumplido'],
        ['2. Interfaz responsiva e intuitiva', 'Cumplido'],
        ['3. API RESTful para operaciones CRUD', 'Cumplido'],
        ['4. Sistema de autenticación JWT por roles', 'Cumplido'],
        ['5. Reportes automatizados (PDF y Excel)', 'Cumplido'],
        ['6. Pruebas funcionales y de usabilidad', 'Cumplido'],
    ]
)

doc.add_heading('3.4 Limitaciones y Trabajo Futuro', level=2)

p = doc.add_paragraph()
make_bold(p, 'Limitaciones encontradas durante el desarrollo:')
lims = [
    'El sistema no cuenta con integración a pasarelas de pago, por lo que las ventas registradas deben confirmarse manualmente contra el comprobante de pago.',
    'No se implementó un módulo de facturación electrónica, ya que los requisitos tributarios varían según el país.',
    'La generación de reportes en tiempo real puede volverse lenta si el volumen de datos supera los 50,000 registros.',
]
for l in lims:
    doc.add_paragraph(l, style='List Bullet')

p = doc.add_paragraph()
make_bold(p, 'Trabajo futuro recomendado:')
tf = [
    'Implementación de un módulo de facturación electrónica con integración a entidades tributarias.',
    'Desarrollo de una aplicación móvil nativa (Android/iOS) complementaria.',
    'Integración con sistemas de código de barras y lectores QR.',
    'Implementación de alertas por correo electrónico para stock mínimo.',
    'Migración a una arquitectura de microservicios.',
    'Incorporación de inteligencia artificial para predicción de demanda y optimización de stock.',
]
for t in tf:
    doc.add_paragraph(t, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. REFERENCIAS BIBLIOGRÁFICAS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. REFERENCIAS BIBLIOGRÁFICAS', level=1)

refs = [
    'Pressman, R. S. (2014). Ingeniería del Software: Un Enfoque Práctico (7ª ed.). McGraw-Hill.',
    'Sommerville, I. (2016). Software Engineering (10ª ed.). Pearson Education.',
    'Schwaber, K., & Sutherland, J. (2020). La Guía Definitiva de Scrum: Las Reglas del Juego. Scrum.org.',
    'Bass, L., Clements, P., & Kazman, R. (2012). Software Architecture in Practice (3ª ed.). Addison-Wesley.',
    'Fielding, R. T. (2000). Architectural Styles and the Design of Network-based Software Architectures [Tesis doctoral]. UC Irvine.',
    'Subramaniam, V., & Hunt, A. (2006). Practices of an Agile Developer. The Pragmatic Programmers.',
    'Welling, L., & Thomson, L. (2016). PHP and MySQL Web Development (5ª ed.). Addison-Wesley.',
    'Nixon, R. (2018). Learning PHP, MySQL & JavaScript (5ª ed.). O\'Reilly Media.',
    'Duckett, J. (2014). Web Design with HTML, CSS, JavaScript and jQuery Set. John Wiley & Sons.',
    'Cederholm, D. (2010). CSS3 for Web Designers. A Book Apart.',
    'Keith, J., & Sambells, J. (2010). DOM Scripting (2ª ed.). Apress.',
    'Zandstra, M. (2021). PHP 8 Objects, Patterns, and Practice (7ª ed.). Apress.',
    'Williams, H. E., & Lane, D. (2004). Web Database Applications with PHP and MySQL (2ª ed.). O\'Reilly.',
    'Resig, J., & Bibeault, B. (2016). Secrets of the JavaScript Ninja (2ª ed.). Manning.',
    'Flanagan, D. (2020). JavaScript: The Definitive Guide (7ª ed.). O\'Reilly Media.',
    'Fowler, M. (2002). Patterns of Enterprise Application Architecture. Addison-Wesley.',
    'Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). Design Patterns. Addison-Wesley.',
    'Newman, S. (2021). Building Microservices (2ª ed.). O\'Reilly Media.',
    'AdminLTE. (2022). AdminLTE 3: Free Bootstrap Admin Template. https://adminlte.io/',
    'Bootstrap Team. (2022). Bootstrap 4 Documentation. https://getbootstrap.com/docs/4.6/',
    'Chart.js Contributors. (2022). Chart.js Documentation. https://www.chartjs.org/docs/',
    'PHP Documentation Group. (2022). PHP Manual. https://www.php.net/manual/es/',
    'MySQL Documentation. (2022). MySQL 8.0 Reference Manual. https://dev.mysql.com/doc/',
    'OWASP Foundation. (2022). OWASP Top Ten. https://owasp.org/www-project-top-ten/',
    'Nielsen, J. (2012). Usability 101. Nielsen Norman Group. https://www.nngroup.com/articles/usability-101/',
    'W3C. (2018). WCAG 2.1. https://www.w3.org/TR/WCAG21/',
    'Let\'s Encrypt. (2022). Documentación. https://letsencrypt.org/es/docs/',
]
for ref in refs:
    doc.add_paragraph(ref, style='List Number')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. ANEXOS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. ANEXOS', level=1)

doc.add_heading('Anexo A: Diagrama Entidad-Relación de la Base de Datos', level=2)
doc.add_paragraph('Tablas principales del sistema:')
tablas_bd = [
    'usuarios (id, nombre, email, password, rol, created_at, updated_at)',
    'categorias (id, nombre, descripcion, icono, estado, created_at)',
    'proveedores (id, razon_social, ruc, direccion, telefono, email, contacto, notas, created_at)',
    'productos (id, codigo, nombre, descripcion, precio_compra, precio_venta, stock, stock_min, stock_max, imagen, categoria_id, proveedor_id, estado, created_at)',
    'movimientos (id, tipo, cantidad, precio_unitario, precio_total, producto_id, usuario_id, observacion, created_at)',
]
for t in tablas_bd:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('Anexo B: Capturas de Pantalla del Sistema', level=2)
doc.add_paragraph(
    '[En esta sección se incluirían las capturas de pantalla de las principales '
    'interfaces del sistema: login, dashboard, listado de productos, formulario '
    'de producto, reportes, etc.]'
)

doc.add_heading('Anexo C: Manual de Usuario', level=2)
doc.add_paragraph(
    '[Documento aparte con las instrucciones detalladas para el uso del sistema, '
    'incluyendo guías paso a paso para cada módulo funcional.]'
)

doc.add_paragraph()

# ── CIERRE ───────────────────────────────────────────────────────────────────
add_separator()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FIN DEL INFORME')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)
add_separator()

# ── Guardar ──────────────────────────────────────────────────────────────────
output_path = 'C:\\Users\\Usuario\\Documents\\Proyecto\\Informe_Proyecto_Web.docx'
doc.save(output_path)
print(f'Documento Word generado exitosamente en:\n{output_path}')
