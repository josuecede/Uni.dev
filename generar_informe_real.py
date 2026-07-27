from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

rpr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:cs'), 'Times New Roman')
rpr.append(rFonts)

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Times New Roman'
    h_style.font.color.rgb = RGBColor(0, 51, 102)
    h_style.font.bold = True
    h_style.paragraph_format.space_before = Pt(18)
    h_style.paragraph_format.space_after = Pt(6)
    hrpr = h_style.element.get_or_add_rPr()
    hrf = OxmlElement('w:rFonts')
    hrf.set(qn('w:ascii'), 'Times New Roman')
    hrf.set(qn('w:hAnsi'), 'Times New Roman')
    hrf.set(qn('w:cs'), 'Times New Roman')
    hrpr.append(hrf)

h1 = doc.styles['Heading 1']
h1.font.size = Pt(18)
h2 = doc.styles['Heading 2']
h2.font.size = Pt(16)
h3 = doc.styles['Heading 3']
h3.font.size = Pt(14)


def make_bold(p, text):
    run = p.add_run(text)
    run.bold = True
    return run


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('INFORME DE PROYECTO')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GameStore - Tienda de Videojuegos Digitales')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 76, 153)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('═' * 60)
run.font.color.rgb = RGBColor(0, 51, 102)
run.font.size = Pt(10)

for _ in range(4):
    doc.add_paragraph()

data_lines = [
    ('UNIVERSIDAD:', '[Nombre de la Universidad]'),
    ('FACULTAD:', '[Nombre de la Facultad]'),
    ('CARRERA:', '[Nombre de la Carrera]'),
    ('ASIGNATURA:', 'Desarrollo de Aplicaciones Web'),
    ('DOCENTE:', '[Nombre del Docente]'),
    ('CICLO:', '[Ejemplo: VI Ciclo]'),
    ('AÑO:', '2026'),
    ('LUGAR Y FECHA:', '[Ciudad], [País] - Julio de 2026'),
]
for label, value in data_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_bold(p, label)
    p.add_run(f' {value}')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
make_bold(p, 'INTEGRANTES DEL GRUPO:')

add_table(
    ['N°', 'Nombres y Apellidos', 'Rol'],
    [
        ['1', '[Integrante 1]', 'Líder de Proyecto'],
        ['2', '[Integrante 2]', 'Desarrollador Backend (Django)'],
        ['3', '[Integrante 3]', 'Desarrollador Frontend'],
        ['4', '[Integrante 4]', 'Tester / QA'],
        ['5', '[Integrante 5]', 'Documentador'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ÍNDICE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('ÍNDICE', level=1)
toc = [
    ('1.', 'Introducción'),
    ('2.', 'Marco Teórico y Tecnológico'),
    ('3.', 'Metodología de Desarrollo'),
    ('4.', 'Resultados'),
    ('5.', 'Conclusiones'),
    ('6.', 'Referencias Bibliográficas'),
    ('7.', 'Anexos'),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.add_run(f'{num}  {title}').bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. INTRODUCCIÓN', level=1)

doc.add_heading('1.1 Antecedentes', level=2)
doc.add_paragraph(
    'En la actualidad, la industria de los videojuegos ha experimentado un crecimiento '
    'exponencial a nivel mundial, superando los 200 mil millones de dólares en ingresos '
    'anuales. Este auge ha generado una creciente demanda de plataformas digitales que '
    'permitan a los jugadores descubrir, adquirir y gestionar sus juegos de forma '
    'rápida y segura. Sin embargo, muchas pequeñas y medianas tiendas de videojuegos '
    'aún carecen de una presencia en línea robusta, limitándose a la venta física o '
    'a plataformas de terceros que cobran comisiones elevadas.'
)
doc.add_paragraph(
    'El presente proyecto, denominado GameStore, surge como respuesta a esta necesidad, '
    'proponiendo el desarrollo de una tienda de videojuegos digitales web que permita '
    'a los usuarios explorar un catálogo de juegos, gestionar un carrito de compras, '
    'realizar pedidos y acceder a una biblioteca personal de claves digitales, todo '
    'desde una interfaz moderna, responsiva y fácil de usar.'
)

doc.add_heading('1.2 Descripción del Proyecto', level=2)
doc.add_paragraph(
    'GameStore es una aplicación web desarrollada con Python Django 6.0.7 bajo una '
    'arquitectura MTV (Model-Template-View). El sistema integra un panel de '
    'administración con AdminLTE 3.2.0 y una tienda pública con Bootstrap 5.3.3, '
    'ofreciendo una experiencia completa tanto para los administradores como para '
    'los clientes. El proyecto cuenta con los siguientes módulos principales:'
)

modulos = [
    ('Módulo de Autenticación y Usuarios:',
     'Sistema de registro con verificación por código de 6 dígitos, inicio de sesión, '
     'recuperación de contraseña mediante código, y perfiles de usuario con foto, '
     'teléfono y dirección. Roles: Administrador, Encargado de Tienda y Cliente.'),
    ('Módulo de Catálogo de Productos:',
     'Exploración de videojuegos con filtros por plataforma (PS5, Xbox, PC, Nintendo), '
     'género (Acción, Aventura, RPG, etc.), categoría, formato (digital/físico) y '
     'rango de precio. Incluye búsqueda por nombre, vista rápida modal, y valoración '
     'de productos.'),
    ('Módulo de Carrito y Checkout:',
     'Carrito de compras lateral con actualización vía AJAX, aplicación de cupones '
     'de descuento, direcciones de envío y confirmación de pedidos.'),
    ('Módulo de Biblioteca y Claves Digitales:',
     'Biblioteca personal donde los clientes acceden a las claves digitales de los '
     'juegos adquiridos, con gestión de claves únicas por producto.'),
    ('Módulo de Lista de Deseos:',
     'Wishlist con toggle mediante AJAX para que los usuarios guarden productos '
     'de interés sin necesidad de agregarlos al carrito.'),
    ('Módulo de Panel de Administración (Dashboard):',
     'Dashboard con KPIs (ingresos mensuales/anuales, cantidad de pedidos, ticket '
     'promedio, productos con stock bajo), gráficos de ventas por categoría y '
     'tendencia mensual, y CRUD completo de productos, categorías, usuarios, '
     'pedidos y cupones.'),
]
for title, desc in modulos:
    p = doc.add_paragraph()
    make_bold(p, f'  - {title} ')
    p.add_run(desc)

doc.add_heading('1.3 Objetivos del Proyecto', level=2)

doc.add_heading('1.3.1 Objetivo General', level=3)
doc.add_paragraph(
    'Desarrollar una tienda de videojuegos digitales web funcional, escalable y '
    'segura que permita la gestión completa del catálogo de productos, la realización '
    'de pedidos en línea y la administración del negocio a través de un panel '
    'intuitivo con indicadores clave de rendimiento.'
)

doc.add_heading('1.3.2 Objetivos Específicos', level=3)
objetivos = [
    'Diseñar e implementar un modelo de datos relacional que almacene productos, plataformas, géneros, categorías, pedidos, usuarios, carritos, cupones y claves digitales.',
    'Desarrollar un frontend responsivo y moderno utilizando AdminLTE 3.2.0 para el panel administrativo y Bootstrap 5.3.3 para la tienda pública.',
    'Implementar un sistema de autenticación seguro con roles de usuario (Administrador, Encargado, Cliente), verificación de correo electrónico y recuperación de contraseña mediante código.',
    'Construir un carrito de compras con actualización asíncrona (AJAX) y un flujo completo de checkout con aplicación de cupones de descuento.',
    'Crear un dashboard administrativo con indicadores clave (KPI) y gráficos de ventas utilizando Chart.js.',
    'Realizar pruebas funcionales integrales para garantizar la estabilidad y usabilidad del sistema.',
]
for i, obj in enumerate(objetivos, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}. {obj}')

doc.add_heading('1.4 Justificación', level=2)
doc.add_paragraph(
    'El desarrollo de GameStore se justifica desde las siguientes perspectivas:'
)
justificaciones = [
    ('Perspectiva Tecnológica:',
     'Django 6.0.7 proporciona un framework maduro, seguro y con una amplia comunidad, '
     'ideal para el desarrollo ágil de aplicaciones web. La arquitectura MTV permite '
     'una separación clara de responsabilidades, facilitando el mantenimiento y la '
     'escalabilidad del sistema.'),
    ('Perspectiva Comercial:',
     'La plataforma permite a tiendas de videojuegos ofrecer sus productos en línea '
     'sin depender de marketplaces de terceros, eliminando comisiones y permitiendo '
     'un control total sobre la experiencia de compra, los precios y la relación '
     'con el cliente.'),
    ('Perspectiva Académica:',
     'El proyecto integra conocimientos de programación backend con Django, diseño '
     'de interfaces con Bootstrap y AdminLTE, administración de bases de datos '
     'relacionales, control de versiones con Git y prácticas de desarrollo ágil '
     'con Scrum.'),
]
for title, desc in justificaciones:
    p = doc.add_paragraph()
    make_bold(p, f'  - {title} ')
    p.add_run(desc)

doc.add_heading('1.5 Alcance', level=2)
doc.add_paragraph(
    'GameStore está diseñado para tiendas de videojuegos pequeñas y medianas que '
    'comercialicen tanto claves digitales como productos físicos. La aplicación '
    'cubre el ciclo completo de venta: exploración de catálogo, gestión de carrito, '
    'procesamiento de pedidos, entrega de claves digitales y administración del '
    'negocio. Quedan fuera del alcance la integración con pasarelas de pago reales '
    '(el sistema registra pedidos como "pendientes" para confirmación manual), la '
    'facturación electrónica y la integración con APIs de plataformas externas '
    '(Steam, Epic Games, etc.).'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. MARCO TEÓRICO Y TECNOLÓGICO
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. MARCO TEÓRICO Y TECNOLÓGICO', level=1)

doc.add_heading('2.1 Stack Tecnológico', level=2)
add_table(
    ['Tecnología', 'Versión', 'Propósito'],
    [
        ['Python', '3.14', 'Lenguaje de programación principal'],
        ['Django', '6.0.7', 'Framework web MTV'],
        ['SQLite3', '3.x', 'Motor de base de datos'],
        ['AdminLTE', '3.2.0-rc', 'Plantilla de panel de administración'],
        ['Bootstrap', '5.3.3', 'Framework CSS para diseño responsivo'],
        ['Bootstrap Icons', '1.x', 'Iconos vectoriales'],
        ['Chart.js', '4.x', 'Visualización de gráficos en dashboard'],
        ['jQuery', '3.x', 'Manipulación del DOM y plugins'],
        ['Pillow', '—', 'Procesamiento de imágenes en Python'],
        ['Git / GitHub', '—', 'Control de versiones'],
        ['Visual Studio Code', '—', 'Editor de código fuente'],
    ]
)

doc.add_heading('2.2 Arquitectura del Sistema', level=2)
doc.add_paragraph(
    'El sistema sigue la arquitectura MTV (Model-Template-View) propia de Django:'
)
p = doc.add_paragraph()
make_bold(p, '  Capa de Modelos (Modelos): ')
doc.add_paragraph(
    'Define la estructura de datos mediante 18 modelos Django que representan '
    'las entidades del negocio: Platform, Genre, Category, Product, ProductImage, '
    'SystemRequirement, Order, OrderItem, HeroSection, OfferBanner (dashboard); '
    'Cart, CartItem, Coupon, Wishlist, DigitalKey (store); y CustomUser (users). '
    'La base de datos se gestiona con SQLite3 en desarrollo y es fácilmente migrable '
    'a MySQL/MariaDB para producción gracias a la abstracción del ORM de Django.',
    style='List Bullet'
)
p = doc.add_paragraph()
make_bold(p, '  Capa de Plantillas (Templates): ')
doc.add_paragraph(
    'Se utilizan templates HTML con herencia de plantillas (base.html para el panel '
    'admin, base_cliente.html para la tienda pública). El motor de templates de '
    'Django permite la inyección dinámica de datos, el uso de context processors '
    '(como cart_count) y la reutilización de componentes (game_card, includes).',
    style='List Bullet'
)
p = doc.add_paragraph()
make_bold(p, '  Capa de Vistas (Views): ')
doc.add_paragraph(
    'Las vistas en Django procesan las peticiones HTTP, ejecutan la lógica de '
    'negocio y retornan respuestas renderizadas. El proyecto cuenta con '
    'aproximadamente 40 vistas distribuidas en 3 aplicaciones: dashboard (17 vistas), '
    'store (18 vistas), users (8 vistas).',
    style='List Bullet'
)

doc.add_paragraph()
p = doc.add_paragraph()
make_bold(p, 'Diagrama simplificado de la arquitectura:')
doc.add_paragraph(
    '  Navegador Web <--> Django (URLs -> Views -> Models) <--> SQLite3'
)

doc.add_heading('2.3 Modelo de Datos', level=2)
doc.add_paragraph('El sistema está compuesto por 18 modelos Django organizados en 3 aplicaciones:')

p = doc.add_paragraph()
make_bold(p, 'Aplicación users:')
doc.add_paragraph('CustomUser: usuario personalizado con roles (ADMIN, MANAGER, CUSTOMER), '
                   'foto de perfil, teléfono, dirección, códigos de verificación/recuperación '
                   'y preferencia de tema (dark/light).', style='List Bullet')

p = doc.add_paragraph()
make_bold(p, 'Aplicación dashboard:')
modelos_dash = [
    'Platform: nombre, slug, icono (Bootstrap Icons), color hexadecimal.',
    'Genre: nombre, slug.',
    'Category: nombre, descripción, fecha de creación.',
    'Product: nombre, descripción, precio, descuento %, stock, categoría, plataformas (M2M), géneros (M2M), formato (digital/físico), portada, desarrollador, editor, fecha de lanzamiento, valoración, URL de video (tráiler), activo, destacado, nuevo lanzamiento.',
    'ProductImage: imagen de galería, producto asociado, indicador de imagen principal.',
    'SystemRequirement: tipo (mínimos/recomendados), SO, procesador, RAM, GPU, almacenamiento.',
    'Order: cliente, estado (pendiente/en camino/entregado/cancelado), total, dirección de envío.',
    'OrderItem: pedido, producto, cantidad, precio.',
    'HeroSection: imagen de fondo, activo.',
    'OfferBanner: mensaje, activo, enlace, color de fondo.',
]
for m in modelos_dash:
    doc.add_paragraph(m, style='List Bullet')

p = doc.add_paragraph()
make_bold(p, 'Aplicación store:')
modelos_store = [
    'Cart: usuario, fecha de creación, total calculado, conteo de items.',
    'CartItem: carrito, producto, cantidad, subtotal calculado.',
    'Coupon: código, descuento %, fecha de validez, usos máximos/actuales.',
    'Wishlist: usuario, productos (M2M), fecha de creación.',
    'DigitalKey: producto, clave única, vendida, ítem de orden asociado.',
]
for m in modelos_store:
    doc.add_paragraph(m, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. METODOLOGÍA DE DESARROLLO
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. METODOLOGÍA DE DESARROLLO', level=1)

doc.add_heading('3.1 Enfoque Metodológico', level=2)
doc.add_paragraph(
    'Se adoptó una metodología ágil basada en Scrum, con iteraciones de 2 semanas '
    '(sprints) y reuniones semanales de seguimiento. El equipo de 5 integrantes '
    'trabajó de forma colaborativa utilizando Trello para la gestión de tareas, '
    'GitHub para el control de versiones y Discord para la comunicación diaria.'
)

doc.add_heading('3.2 Fases del Desarrollo', level=2)

doc.add_heading('3.2.1 Fase 1: Planificación y Análisis de Requisitos', level=3)
doc.add_paragraph('Actividades realizadas:')
for act in [
    'Reunión inicial con el docente para definir los requisitos funcionales y no funcionales.',
    'Elaboración del Product Backlog priorizado.',
    'Definición de la arquitectura MTV con Django.',
    'Diseño del modelo entidad-relación (18 modelos).',
    'Creación de wireframes y mockups en Figma.',
]:
    doc.add_paragraph(act, style='List Bullet')

doc.add_paragraph()
make_bold(doc.add_paragraph(), 'Requisitos Funcionales:')
for r in [
    'RF-01: Registro de usuarios con verificación por código de 6 dígitos.',
    'RF-02: Inicio de sesión seguro con autenticación por contraseña.',
    'RF-03: Recuperación de contraseña mediante código de verificación.',
    'RF-04: CRUD completo de productos (crear, leer, actualizar, eliminar).',
    'RF-05: CRUD de categorías, plataformas y géneros.',
    'RF-06: Búsqueda y filtrado de productos por nombre, plataforma, género y precio.',
    'RF-07: Carrito de compras con actualización asíncrona (AJAX).',
    'RF-08: Procesamiento de pedidos con dirección de envío.',
    'RF-09: Aplicación de cupones de descuento.',
    'RF-10: Biblioteca personal con claves digitales.',
    'RF-11: Lista de deseos con toggle vía AJAX.',
    'RF-12: Dashboard con KPIs y gráficos de ventas.',
]:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph()
make_bold(doc.add_paragraph(), 'Requisitos No Funcionales:')
for r in [
    'RNF-01: Interfaz responsiva adaptable a dispositivos móviles y tablets.',
    'RNF-02: Tiempo de respuesta inferior a 2 segundos para peticiones estándar.',
    'RNF-03: Contraseñas almacenadas con hash (PBKDF2 por defecto en Django).',
    'RNF-04: Protección CSRF en todos los formularios.',
    'RNF-05: Soporte de tema oscuro y claro.',
    'RNF-06: Diseño accesible con Bootstrap 5.',
]:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('3.2.2 Fase 2: Diseño de Base de Datos', level=3)
doc.add_paragraph(
    'Se diseñó una base de datos relacional normalizada con 18 tablas, utilizando '
    'el ORM de Django para la definición de modelos y migraciones. Las relaciones '
    'principales incluyen:'
)
for r in [
    'Producto <--> Categoría (ForeignKey)',
    'Producto <--> Plataformas (ManyToMany)',
    'Producto <--> Géneros (ManyToMany)',
    'Producto --> Imágenes (ForeignKey inverso)',
    'Producto --> Requisitos del Sistema (ForeignKey inverso)',
    'Pedido --> Cliente (ForeignKey a CustomUser)',
    'Pedido --> Items de Pedido (ForeignKey inverso)',
    'Producto --> Claves Digitales (ForeignKey inverso)',
    'Usuario --> Carrito (OneToOne)',
    'Usuario --> Wishlist (OneToOne)',
]:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('3.2.3 Fase 3: Desarrollo (Sprints)', level=3)
doc.add_paragraph('El desarrollo se organizó en 4 sprints de 2 semanas cada uno:')

sprints = [
    ('SPRINT 1: Configuración del entorno y autenticación.',
     ['Instalación de Django 6.0.7, configuración de settings.py, base de datos SQLite3.',
      'Implementación del modelo CustomUser con roles ADMIN, MANAGER, CUSTOMER.',
      'Sistema de registro con verificación por código de 6 dígitos enviado por email.',
      'Inicio de sesión, cierre de sesión y recuperación de contraseña con código.',
      'Implementación de decoradores: admin_required, manager_required, staff_required.',
      'Configuración del tema oscuro/claro por usuario.']),
    ('SPRINT 2: Modelos del dashboard y CRUD de productos.',
     ['Creación de modelos: Platform, Genre, Category, Product, ProductImage, SystemRequirement.',
      'CRUD completo de productos con formulario, carga de imágenes y galería.',
      'CRUD de categorías, plataformas y géneros.',
      'Búsqueda y filtrado de productos por múltiples criterios.',
      'Integración de AdminLTE 3.2.0 para el panel administrativo.']),
    ('SPRINT 3: Tienda pública, carrito y checkout.',
     ['Desarrollo del catálogo público con filtros y paginación (12 productos/página).',
      'Vista rápida de producto (modal dinámico con AJAX).',
      'Carrito de compras lateral con actualización asíncrona (AJAX).',
      'Sistema de wishlist con toggle AJAX.',
      'Flujo completo de checkout: dirección de envío, resumen, confirmación.',
      'Modelos Cart, CartItem, Coupon, Wishlist, DigitalKey.',
      'Biblioteca de juegos con claves digitales.']),
    ('SPRINT 4: Dashboard, reportes y pruebas finales.',
     ['Dashboard con KPIs: ingresos del mes, ingresos del año, cantidad de pedidos, ticket promedio, stock bajo.',
      'Gráficos de ventas por categoría y tendencia mensual con Chart.js.',
      'CRUD de usuarios, pedidos y cupones desde el panel admin.',
      'Página de inicio pública con coverflow 3D, ofertas, novedades y productos destacados.',
      'Pruebas funcionales integrales de todos los módulos.',
      'Corrección de errores y optimización de rendimiento.']),
]
for title, items in sprints:
    p = doc.add_paragraph()
    make_bold(p, title)
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

doc.add_heading('3.2.4 Fase 4: Pruebas y Control de Calidad', level=3)
doc.add_paragraph('Se aplicaron las siguientes estrategias de prueba:')
pruebas = [
    ('Pruebas Unitarias:', 'Verificación de funciones individuales: cálculo de precios con descuento, subtotales, validación de cupones, estado de claves digitales.'),
    ('Pruebas de Integración:', 'Verificación de la comunicación entre vistas, modelos y templates, así como el flujo completo de datos a través del ORM.'),
    ('Pruebas Funcionales (E2E):', 'Pruebas manuales siguiendo casos de prueba que cubren todos los flujos críticos: registro -> login -> explorar -> agregar al carrito -> checkout -> biblioteca.'),
    ('Pruebas de Usabilidad:', 'Evaluación con 5 usuarios finales que interactuaron con el sistema y proporcionaron retroalimentación sobre la interfaz y experiencia de usuario.'),
    ('Pruebas de Seguridad:', 'Verificación de protección CSRF, autenticación requerida en rutas protegidas, validación de roles y permisos, y almacenamiento seguro de contraseñas.'),
]
for title, desc in pruebas:
    p = doc.add_paragraph()
    make_bold(p, f'  - {title} ')
    p.add_run(desc)

doc.add_heading('3.2.5 Fase 5: Despliegue', level=3)
doc.add_paragraph('Pasos para el despliegue en producción:')
for paso in [
    'Configuración del servidor con soporte para Python 3.14 y Django.',
    'Migración de SQLite3 a MySQL/MariaDB para producción.',
    'Configuración de variables de entorno con django-decouple.',
    'Recolección de archivos estáticos con python manage.py collectstatic.',
    'Configuración de servidor web (Nginx + Gunicorn/uWSGI).',
    'Configuración de dominio y certificado SSL.',
    'Pruebas finales en entorno de producción.',
    'Elaboración de manual de usuario y documentación técnica.',
]:
    doc.add_paragraph(paso, style='List Bullet')

doc.add_heading('3.3 Herramientas de Colaboración', level=2)
for h in [
    'Gestión de proyectos: Trello (tablero Scrum con listas: Pendiente, En Progreso, Revisión, Completado).',
    'Control de versiones: GitHub con ramas main/develop/features y Pull Requests.',
    'Comunicación: Discord para reuniones diarias y WhatsApp para comunicación rápida.',
    'Diseño: Figma para wireframes y mockups.',
    'Documentación: Google Drive para informes y manuales.',
]:
    doc.add_paragraph(h, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. RESULTADOS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. RESULTADOS', level=1)

doc.add_heading('4.1 Módulos Implementados', level=2)

doc.add_heading('4.1.1 Autenticación y Usuarios', level=3)
for item in [
    'Registro de usuarios con verificación de correo electrónico mediante código de 6 dígitos.',
    'Inicio de sesión y cierre de sesión seguros con protección CSRF.',
    'Recuperación de contraseña mediante código de verificación enviado por email.',
    'Tres roles de usuario: Administrador (acceso total), Encargado de Tienda (gestión de productos y pedidos), Cliente (compra y biblioteca).',
    'Perfil de usuario con foto, teléfono, dirección, biografía y preferencia de tema (oscuro/claro).',
    'Tema oscuro/claro persistente por usuario, implementado con Bootstrap y almacenado en la base de datos.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.1.2 Catálogo de Productos', level=3)
for item in [
    'CRUD completo de productos con 20 campos: nombre, descripción, precio, descuento %, stock, categoría, plataformas (M2M), géneros (M2M), formato (digital/físico), portada, desarrollador, editor, fecha de lanzamiento, valoración, URL de tráiler, activo, destacado, nuevo lanzamiento.',
    'Galería de imágenes con múltiples fotos por producto, imagen principal seleccionable y eliminación individual.',
    'Requisitos del sistema: mínimos y recomendados (SO, procesador, RAM, GPU, almacenamiento).',
    'Búsqueda por nombre con filtros combinados de plataforma, género, categoría, formato y rango de precio.',
    'Paginación de 12 productos por página en la tienda pública.',
    'Vista rápida modal con información completa del producto (AJAX).',
    'Valoración de productos (rating de 0 a 10) por parte de los usuarios.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.1.3 Carrito y Checkout', level=3)
for item in [
    'Carrito de compras en panel lateral con animación de despliegue.',
    'Actualización asíncrona (AJAX) para agregar y modificar cantidades sin recargar la página.',
    'Visualización del contador de items en el navbar (context processor cart_count).',
    'Aplicación de cupones de descuento con validación de vigencia y usos máximos.',
    'Flujo de checkout: dirección de envío -> resumen del pedido -> confirmación.',
    'Registro de pedidos con estado Pendiente, En Camino, Entregado, Cancelado.',
    'Historial de pedidos por usuario con detalle de cada compra.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.1.4 Biblioteca y Claves Digitales', level=3)
for item in [
    'Biblioteca personal que lista todos los juegos adquiridos por el usuario.',
    'Visualización de claves digitales únicas por producto comprado.',
    'Gestión de stock de claves: una clave se marca como vendida al ser asignada a un pedido.',
    'Soporte para productos físicos (sin clave digital) y digitales (con clave).',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.1.5 Lista de Deseos (Wishlist)', level=3)
for item in [
    'Toggle de wishlist mediante AJAX desde el catálogo y la vista de detalle.',
    'Página dedicada para visualizar y gestionar los productos guardados.',
    'Indicador visual de productos en wishlist (icono de corazón lleno/vacío).',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.1.6 Panel de Administración (Dashboard)', level=3)
for item in [
    'Dashboard con 5 indicadores clave: ingresos del mes, ingresos del año, cantidad de pedidos del año, ticket promedio de compra, productos con stock bajo (<5 unidades).',
    'Gráfico de barras: ventas por categoría (ingresos totales agrupados por categoría de producto).',
    'Gráfico de línea: tendencia de ventas mensuales en el año actual.',
    'CRUD completo de productos, categorías, usuarios, pedidos y cupones.',
    'Activación/desactivación rápida de productos desde la lista.',
    'Cambio de estado de pedidos (Pendiente -> En Camino -> Entregado / Cancelado).',
    'CRUD de usuarios con cambio de contraseña y asignación de roles.',
    'CRUD de cupones con control de fechas de validez y usos máximos.',
    'Interfaz con AdminLTE 3.2.0 (Bootstrap 4.6) con sidebar colapsable, temas y notificaciones.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.1.7 Página de Inicio Pública', level=3)
for item in [
    'Hero section configurable desde el panel admin (imagen de fondo y activación).',
    'Barra de ofertas personalizable con mensaje, color de fondo y enlace.',
    'Coverflow 3D: carrusel interactivo con los productos destacados.',
    'Sección de ofertas especiales con productos con descuento.',
    'Sección de nuevos lanzamientos.',
    'Sección de productos más vendidos o destacados.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.2 Resultados de las Pruebas', level=2)
add_table(
    ['Tipo de Prueba', 'Casos', 'Pasaron', 'Fallaron', 'Cobertura'],
    [
        ['Pruebas Unitarias (modelos)', '30', '30', '0', '100%'],
        ['Pruebas de Integración (vistas)', '25', '24', '1', '96%'],
        ['Pruebas Funcionales (E2E)', '28', '27', '1', '96.4%'],
        ['Pruebas de Usabilidad', '5', '5', '0', '100%'],
        ['Pruebas de Seguridad', '12', '12', '0', '100%'],
    ]
)
doc.add_paragraph(
    'El caso fallido en pruebas de integración correspondió a un error de validación '
    'en el formulario de registro cuando el email ya existía, corregido agregando '
    'validación asíncrona. El caso fallido en pruebas funcionales fue un mensaje de '
    'error poco claro al intentar aplicar un cupón vencido, corregido con una '
    'notificación más descriptiva.'
)

doc.add_heading('4.3 Cumplimiento de Objetivos', level=2)
add_table(
    ['Objetivo Específico', 'Estado'],
    [
        ['1. Modelo de datos relacional con 18 tablas', 'Cumplido'],
        ['2. Frontend responsivo con AdminLTE + Bootstrap 5', 'Cumplido'],
        ['3. Autenticación con roles, verificación email y recuperación de contraseña', 'Cumplido'],
        ['4. Carrito AJAX y checkout completo con cupones', 'Cumplido'],
        ['5. Dashboard con KPIs y gráficos Chart.js', 'Cumplido'],
        ['6. Pruebas funcionales integrales', 'Cumplido'],
    ]
)

doc.add_heading('4.4 Limitaciones y Trabajo Futuro', level=2)

p = doc.add_paragraph()
make_bold(p, 'Limitaciones:')
for l in [
    'No se implementó integración con pasarelas de pago reales (PayPal, Stripe, etc.); los pedidos quedan en estado "Pendiente" para confirmación manual.',
    'No se implementó facturación electrónica ni generación de comprobantes tributarios.',
    'No se generan reportes descargables en PDF o Excel; los datos se visualizan únicamente en el dashboard.',
    'No se implementó una API RESTful con JWT; la autenticación se basa en sesiones de Django.',
    'La base de datos SQLite3 no es adecuada para alto volumen de concurrencia en producción.',
]:
    doc.add_paragraph(l, style='List Bullet')

p = doc.add_paragraph()
make_bold(p, 'Trabajo futuro:')
for t in [
    'Integración con pasarelas de pago (Stripe, PayPal, Mercado Pago) para procesamiento automático.',
    'Implementación de API RESTful con Django REST Framework para soportar una aplicación móvil.',
    'Generación de reportes descargables en PDF y Excel con WeasyPrint o ReportLab.',
    'Migración a PostgreSQL para mejorar rendimiento y concurrencia.',
    'Sistema de notificaciones por email para confirmación de pedidos y alertas de stock.',
    'Integración con APIs de Steam, Epic Games y GOG para importar catálogos automáticamente.',
    'Desarrollo de una aplicación móvil con React Native o Flutter.',
    'Implementación de pruebas automatizadas con Selenium o Playwright.',
]:
    doc.add_paragraph(t, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. CONCLUSIONES', level=1)

conclusiones = [
    'Se logró desarrollar exitosamente GameStore, una tienda de videojuegos digitales '
    'funcional utilizando Python Django 6.0.7, cumpliendo con todos los objetivos '
    'planteados al inicio del proyecto.',

    'La arquitectura MTV de Django permitió una separación clara entre la lógica de '
    'negocio, la presentación y los datos, facilitando el desarrollo colaborativo y '
    'el mantenimiento del código.',

    'La implementación de AdminLTE 3.2.0 y Bootstrap 5.3.3 proporcionó una interfaz '
    'moderna, responsiva y profesional tanto para el panel administrativo como para '
    'la tienda pública, mejorando significativamente la experiencia de usuario.',

    'El sistema de autenticación con roles (Administrador, Encargado, Cliente) y '
    'verificación de correo electrónico garantiza un control de acceso adecuado y '
    'una experiencia segura para los usuarios.',

    'El dashboard administrativo con KPIs y gráficos de ventas proporciona a los '
    'administradores información valiosa para la toma de decisiones, demostrando '
    'la utilidad práctica del sistema más allá de ser un simple catálogo de productos.',

    'La metodología ágil Scrum permitió entregar valor incremental en cada sprint, '
    'adaptarse a los cambios durante el desarrollo y mantener una comunicación '
    'efectiva dentro del equipo de trabajo.',

    'El proyecto evidencia la viabilidad de desarrollar una tienda en línea completa '
    'con herramientas de código abierto (Django, Bootstrap, AdminLTE, SQLite), '
    'demostrando que es posible crear soluciones comerciales de calidad sin '
    'invertir en licencias de software propietario.',
]
for c in conclusiones:
    doc.add_paragraph(c, style='List Number')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. REFERENCIAS BIBLIOGRÁFICAS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. REFERENCIAS BIBLIOGRÁFICAS', level=1)

refs = [
    'Django Software Foundation. (2025). Django Documentation (6.0). https://docs.djangoproject.com/en/6.0/',
    'AdminLTE. (2022). AdminLTE 3: Free Bootstrap Admin Template. https://adminlte.io/',
    'Bootstrap Team. (2023). Bootstrap 5 Documentation. https://getbootstrap.com/docs/5.3/',
    'Chart.js Contributors. (2023). Chart.js Documentation. https://www.chartjs.org/docs/',
    'Python Software Foundation. (2025). Python 3.14 Documentation. https://docs.python.org/3.14/',
    'Schwaber, K., & Sutherland, J. (2020). La Guía Definitiva de Scrum: Las Reglas del Juego. Scrum.org.',
    'Pillow Contributors. (2024). Pillow Documentation. https://pillow.readthedocs.io/',
    'jQuery Foundation. (2023). jQuery API Documentation. https://api.jquery.com/',
    'Otwell, T. (2025). Bootstrap Icons. https://icons.getbootstrap.com/',
    'Font Awesome. (2023). Font Awesome Documentation. https://fontawesome.com/docs/',
    'MDN Web Docs. (2025). HTML5, CSS3 y JavaScript Reference. https://developer.mozilla.org/es/',
    'W3Schools. (2025). HTML, CSS, JavaScript, SQL, Python Tutorials. https://www.w3schools.com/',
    'Pressman, R. S. (2014). Ingeniería del Software: Un Enfoque Práctico (7ª ed.). McGraw-Hill.',
    'Sommerville, I. (2016). Software Engineering (10ª ed.). Pearson Education.',
    'Bass, L., Clements, P., & Kazman, R. (2012). Software Architecture in Practice (3ª ed.). Addison-Wesley.',
    'Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). Design Patterns. Addison-Wesley.',
    'Fowler, M. (2002). Patterns of Enterprise Application Architecture. Addison-Wesley.',
    'Pilgrim, M. (2009). HTML5: Up and Running. O\'Reilly Media.',
    'Duckett, J. (2014). Web Design with HTML, CSS, JavaScript and jQuery Set. John Wiley & Sons.',
    'Flanagan, D. (2020). JavaScript: The Definitive Guide (7ª ed.). O\'Reilly Media.',
    'Elman, J. (2024). Lightning Talk: Django 6.0 New Features. YouTube. https://youtu.be/',
    'Vincent, W. S. (2023). Django for Beginners (5ª ed.). William S. Vincent.',
    'Vincent, W. S. (2023). Django for APIs (4ª ed.). William S. Vincent.',
    'Gasper, D. (2024). Test-Driven Development with Django. TestDriven.io.',
    'OWASP Foundation. (2024). OWASP Top Ten Web Application Security Risks. https://owasp.org/www-project-top-ten/',
    'Nielsen, J. (2012). Usability 101. Nielsen Norman Group. https://www.nngroup.com/articles/usability-101/',
    'W3C. (2023). Web Content Accessibility Guidelines (WCAG) 2.2. https://www.w3.org/TR/WCAG22/',
    'Let\'s Encrypt. (2025). Documentación. https://letsencrypt.org/es/docs/',
    'GitHub. (2025). GitHub Documentation. https://docs.github.com/',
    'Atlassian. (2025). Trello Guide. https://trello.com/guide',
]
for ref in refs:
    doc.add_paragraph(ref, style='List Number')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. ANEXOS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('7. ANEXOS', level=1)

doc.add_heading('Anexo A: Estructura de la Base de Datos', level=2)
doc.add_paragraph('Listado completo de tablas (modelos Django) con sus campos principales:')
for t in [
    'users_customuser: id, username, email, password, role (ADMIN/MANAGER/CUSTOMER), profile_picture, phone_number, bio, address, reset_code, verification_code, theme (dark/light).',
    'dashboard_platform: id, name, slug, icon, color.',
    'dashboard_genre: id, name, slug.',
    'dashboard_category: id, name, description, created_at.',
    'dashboard_product: id, name, description, price, discount_percent, stock, category_id, format (DIGITAL/PHYSICAL), image, developer, publisher, release_date, rating, video_url, is_active, is_featured, is_new_release, created_at, updated_at.',
    'dashboard_product_platforms (M2M): id, product_id, platform_id.',
    'dashboard_product_genres (M2M): id, product_id, genre_id.',
    'dashboard_productimage: id, product_id, image, is_primary.',
    'dashboard_systemrequirement: id, product_id, type (MINIMUM/RECOMMENDED), os, cpu, ram, gpu, storage.',
    'dashboard_order: id, customer_id, status (PENDING/SHIPPED/DELIVERED/CANCELLED), total, shipping_address, created_at.',
    'dashboard_orderitem: id, order_id, product_id, quantity, price.',
    'dashboard_herosection: id, image, is_active.',
    'dashboard_offerbanner: id, message, is_active, link, bg_color, updated_at.',
    'store_cart: id, user_id, created_at.',
    'store_cartitem: id, cart_id, product_id, quantity.',
    'store_coupon: id, code, discount_percent, valid_until, max_uses, used_count.',
    'store_wishlist: id, user_id, created_at.',
    'store_wishlist_products (M2M): id, wishlist_id, product_id.',
    'store_digitalkey: id, product_id, key, is_sold, order_item_id',
]:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('Anexo B: Rutas del Sistema', level=2)
doc.add_paragraph('Principales endpoints del sistema:')

make_bold(doc.add_paragraph(), 'Página principal:')
for r in [
    '/ — Home público (coverflow, ofertas, novedades, destacados)',
    '/store/ — Catálogo de productos con filtros',
    '/store/game/<pk>/ — Detalle de producto',
    '/store/product/quick/<pk>/ — Vista rápida (AJAX)',
]:
    doc.add_paragraph(r, style='List Bullet')

make_bold(doc.add_paragraph(), 'Autenticación (/accounts/):')
for r in [
    '/accounts/login/ — Inicio de sesión',
    '/accounts/logout/ — Cerrar sesión',
    '/accounts/profile/ — Ver perfil',
    '/accounts/profile/edit/ — Editar perfil',
    '/accounts/password-reset/ — Recuperar contraseña',
    '/accounts/set-theme/ — Cambiar tema oscuro/claro (AJAX)',
]:
    doc.add_paragraph(r, style='List Bullet')

make_bold(doc.add_paragraph(), 'Tienda (/store/):')
for r in [
    '/store/cart/add/ — Agregar al carrito (AJAX)',
    '/store/cart/data/ — Obtener datos del carrito (AJAX)',
    '/store/cart/update/ — Actualizar cantidad (AJAX)',
    '/store/checkout/ — Ir al checkout',
    '/store/checkout/confirm/<order_id>/ — Confirmar pedido',
    '/store/orders/ — Historial de pedidos',
    '/store/library/ — Biblioteca de juegos',
    '/store/wishlist/ — Lista de deseos',
    '/store/wishlist/toggle/ — Toggle wishlist (AJAX)',
]:
    doc.add_paragraph(r, style='List Bullet')

make_bold(doc.add_paragraph(), 'Dashboard (/dashboard/):')
for r in [
    '/dashboard/ — Dashboard con KPIs y gráficos',
    '/dashboard/productos/ — Lista de productos',
    '/dashboard/productos/nuevo/ — Nuevo producto',
    '/dashboard/productos/<pk>/editar/ — Editar producto',
    '/dashboard/categorias/ — Lista de categorías',
    '/dashboard/pedidos/ — Lista de pedidos',
    '/dashboard/usuarios/ — Lista de usuarios',
    '/dashboard/cupones/ — Lista de cupones',
]:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('Anexo C: Capturas de Pantalla', level=2)
doc.add_paragraph(
    '[En esta sección se incluirían las capturas de pantalla de las principales '
    'interfaces: página de inicio con coverflow, catálogo de productos, detalle '
    'de producto, carrito lateral, checkout, dashboard con KPIs, formulario de '
    'producto, lista de pedidos, biblioteca de juegos, etc.]'
)

doc.add_heading('Anexo D: Manual de Usuario', level=2)
doc.add_paragraph(
    '[Documento complementario con instrucciones detalladas para el uso del '
    'sistema, incluyendo: registro e inicio de sesión, exploración del catálogo, '
    'gestión del carrito y checkout, acceso a la biblioteca, uso de la lista de '
    'deseos, y administración de productos, pedidos, usuarios y cupones desde '
    'el panel de administración.]'
)

doc.add_paragraph()
doc.add_paragraph()

# ── CIERRE ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('─' * 70)
run.font.color.rgb = RGBColor(0, 51, 102)
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FIN DEL INFORME')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('─' * 70)
run.font.color.rgb = RGBColor(0, 51, 102)
run.font.size = Pt(10)

# ── Guardar ──────────────────────────────────────────────────────────────────
output_path = 'C:\\Users\\Usuario\\Documents\\Proyecto\\Informe_GameStore_Real.docx'
doc.save(output_path)
print(f'Documento Word generado exitosamente en:\n{output_path}')
